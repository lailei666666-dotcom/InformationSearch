from __future__ import annotations

import json
from dataclasses import dataclass
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


ROOT = Path(__file__).resolve().parents[2]
TARGET_DOCX = ROOT / "面向用户需求理解的商品评论语义检索增强研究.docx"
BENCHMARK_DIR = ROOT / "outputs" / "runs" / "xiaomi_bailian_benchmark"


@dataclass(slots=True)
class BenchmarkGroup:
    precision: float
    recall: float
    mrr: float


def load_inputs(root: Path) -> dict[str, object]:
    reviews = pd.read_csv(root / "data" / "processed" / "xiaomi_reviews_retrieval.csv")
    queries = pd.read_csv(root / "data" / "annotations" / "queries.csv")
    review_qrels = pd.read_csv(root / "data" / "annotations" / "qrels_reviews.csv")
    product_qrels = pd.read_csv(root / "data" / "annotations" / "qrels_products.csv")
    summary = pd.read_csv(BENCHMARK_DIR / "tables" / "benchmark_summary.csv")
    raw_summary = json.loads((root / "data" / "raw" / "xiaomi_reviews_summary.json").read_text(encoding="utf-8"))

    review_groups = load_query_type_metrics(queries, target="reviews")
    product_groups = load_query_type_metrics(queries, target="products")

    return {
        "reviews": reviews,
        "queries": queries,
        "review_qrels": review_qrels,
        "product_qrels": product_qrels,
        "summary": summary,
        "raw_summary": raw_summary,
        "review_groups": review_groups,
        "product_groups": product_groups,
        "benchmark_dir": BENCHMARK_DIR,
    }


def load_query_type_metrics(
    queries: pd.DataFrame,
    *,
    target: str,
) -> dict[str, dict[str, BenchmarkGroup]]:
    result: dict[str, dict[str, BenchmarkGroup]] = {}
    for system in ("bm25", "semantic", "hybrid"):
        payload = json.loads(
            (BENCHMARK_DIR / "benchmarks" / f"{system}_{target}.json").read_text(
                encoding="utf-8"
            )
        )
        per_query = pd.DataFrame(payload["per_query"])
        merged = queries.merge(per_query, on="query_id")
        grouped = (
            merged.groupby("query_type")[["precision_at_k", "recall_at_k", "reciprocal_rank"]]
            .mean()
            .round(4)
        )
        result[system] = {
            str(query_type): BenchmarkGroup(
                precision=float(row["precision_at_k"]),
                recall=float(row["recall_at_k"]),
                mrr=float(row["reciprocal_rank"]),
            )
            for query_type, row in grouped.iterrows()
        }
    return result


def build_document(output_path: Path) -> None:
    inputs = load_inputs(ROOT)
    doc = Document()
    configure_document(doc)

    add_cover(doc)
    add_abstract(doc, inputs)
    add_section_intro(doc, inputs)
    add_section_theory(doc)
    add_section_system(doc, inputs)
    add_section_experiment(doc, inputs)
    add_section_error_analysis(doc, inputs)
    add_section_conclusion(doc)
    add_references(doc)
    add_acknowledgement(doc)
    add_appendix(doc)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    temp_output = output_path.with_suffix(".tmp.docx")
    doc.save(temp_output)
    temp_output.replace(output_path)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)

    heading1 = doc.styles["Heading 1"]
    heading1.font.name = "黑体"
    heading1._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    heading1.font.size = Pt(16)
    heading1.font.bold = True

    heading2 = doc.styles["Heading 2"]
    heading2.font.name = "黑体"
    heading2._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    heading2.font.size = Pt(14)
    heading2.font.bold = True

    heading3 = doc.styles["Heading 3"]
    heading3.font.name = "黑体"
    heading3._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    heading3.font.size = Pt(12)
    heading3.font.bold = True


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("学校代码：10022")
    set_font(run, "宋体", 12)

    add_spacer(doc, 2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("信息服务与信息检索")
    set_font(run, "黑体", 18, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("2026 年 4 月 27 日")
    set_font(run, "宋体", 12)

    add_spacer(doc, 3)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("面向用户需求理解的商品评论语义检索增强研究")
    set_font(run, "黑体", 20, bold=True)

    add_spacer(doc, 3)

    for line in [
        "作者：赖磊",
        "学号：220401501",
        "专业：信息管理与信息系统",
        "课程：信息服务与检索",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        set_font(run, "宋体", 14)

    doc.add_page_break()


def add_abstract(doc: Document, inputs: dict[str, object]) -> None:
    reviews = inputs["reviews"]
    queries = inputs["queries"]
    summary = inputs["summary"]
    raw_summary = inputs["raw_summary"]
    rows = benchmark_rows(summary)

    doc.add_heading("摘要", level=1)
    abstract_paragraphs = [
        (
            "随着电商平台商品信息和用户评论规模持续增长，用户在检索商品时越来越依赖评论反馈判断商品是否满足"
            "真实使用需求。传统商品检索主要依赖标题、参数和关键词匹配，对“静音键盘”“护眼台灯”这类显式查询具有较好效果，"
            "但当用户使用“宿舍夜里写论文不吵室友的键盘”“出差带着跑客户希望电脑轻一点续航久一点”这类场景化自然语言表达时，"
            "评论反馈中的真实证据往往与查询词不存在直接重合，导致传统检索容易出现漏检、误检和解释性不足。"
        ),
        (
            "本文围绕用户需求理解，设计并实现了一套商品评论语义检索增强平台，统一完成数据采集、评论清洗、BM25 传统检索、"
            "类别约束语义检索、融合检索和指标评测。实验数据来自小米商城公开评价页，覆盖键盘、台灯、耳机、笔记本和充电宝 "
            f"5 个固定品类，共抓取原始评论 {raw_summary['collected_total']} 条，清洗后保留 {len(reviews)} 条检索用评论，"
            f"涉及 {reviews['product_id'].nunique()} 个商品；同时构建了 {len(queries)} 条人工标注查询、"
            f"{len(inputs['review_qrels'])} 条评论级相关标注和 {len(inputs['product_qrels'])} 条商品级相关标注。"
        ),
        (
            "系统实现上，本文以评论为基本检索单元，先利用规则词典识别查询中的商品类别与需求表达，再分别构建 BM25 索引、"
            "FAISS 向量索引和线性融合排序模块，最终以“商品结果 + 评论证据”的方式返回结果。为兼顾课程实验的可复现性和后续升级能力，"
            "平台同时提供本地语义编码器与 OpenAI-compatible embedding 接口，支持在不改动 benchmark 流程的前提下切换到百炼等兼容嵌入服务。"
        ),
        (
            f"实验结果显示，在评论级评测上，BM25 基线取得较高的 Precision@10={rows['bm25_reviews']['precision_at_k']:.4f}、"
            f"Recall@10={rows['bm25_reviews']['recall_at_k']:.4f} 和 MRR={rows['bm25_reviews']['mrr']:.4f}，"
            "说明真实评论场景中存在大量词项重合和评论噪声，传统检索仍具有竞争力；在商品级评测上，融合检索取得最优的 "
            f"MRR={rows['hybrid_products']['mrr']:.4f} 和 Recall@10={rows['hybrid_products']['recall_at_k']:.4f}，"
            "尤其在场景化隐含需求查询中显著优于纯 BM25，表明“类别约束 + 语义检索 + 商品聚合”"
            "更适合最终商品决策场景。本文验证了商品评论语义检索增强系统的工程可行性，也为后续接入商用 embedding、扩展更多品类和"
            "完善商品排序机制提供了可复用的实验平台。"
        ),
    ]
    for text in abstract_paragraphs:
        add_body_paragraph(doc, text)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("关键词：")
    set_font(run, "黑体", 12, bold=True)
    run = p.add_run("用户需求理解；商品评论；BM25；语义检索；FAISS；融合检索；信息检索实验")
    set_font(run, "宋体", 12)


def add_section_intro(doc: Document, inputs: dict[str, object]) -> None:
    doc.add_heading("一 绪论", level=1)
    doc.add_heading("1.1 研究背景", level=2)
    for text in [
        "商品评论已经成为电商检索和消费决策中最有价值的非结构化信息之一。相较于标题、价格和参数，评论能够更直接地反映商品在噪音、续航、重量、佩戴舒适度、宿舍使用、图书馆使用等真实场景下的表现，因此用户越来越倾向于通过评论寻找“是否适合自己”的证据。",
        "然而，用户的商品需求表达通常是自然语言化和场景化的。例如“宿舍晚上用会不会打扰室友的灯”“适合学生党买的耳机”“周末露营两天想带一个能给手机和耳机补电的充电宝”，这类查询并非寻找一个精确参数，而是在寻找与真实场景高度相关的经验反馈。传统基于关键词重合的检索系统很难在评论中稳定发现这些隐含需求。",
        "因此，围绕“类别识别是否正确”“评论反馈能否真正支撑商品推荐”“检索结果能否让用户理解推荐原因”三个问题，构建一个面向用户需求理解的评论反馈检索系统具有明确的课程实践价值和信息服务意义。",
    ]:
        add_body_paragraph(doc, text)

    doc.add_heading("1.2 研究目的与意义", level=2)
    for text in [
        "本文的直接目标不是只做一个语义检索 demo，而是搭建一套完整的实验平台：统一完成数据采集、清洗、标注、三类检索系统构建、评测和结果出图，以便对传统检索、语义检索和融合检索进行可复现对比。",
        "从理论角度看，本文将信息检索课程中的信息采集、索引构建、相关性匹配、排序与评测等经典流程，与向量表示和向量索引技术结合起来，展示了检索系统从关键词匹配向需求理解演进的实际工程路径。",
        "从应用角度看，商品评论语义检索可以将“商品列表”升级为“商品列表 + 评论证据”，帮助用户在购买前快速定位对自己真正有用的反馈，也为智能客服、评论摘要和推荐系统提供更强的证据基础。",
    ]:
        add_body_paragraph(doc, text)

    doc.add_heading("1.3 研究内容与论文结构", level=2)
    for text in [
        "本文的核心工作包括四部分：第一，围绕 5 个固定品类采集并清洗真实评论数据，构建统一实验数据集；第二，实现 BM25 传统检索、类别约束语义检索和融合检索三套系统；第三，建立评论级与商品级双层评测机制，并构建人工标注查询集；第四，将系统实现、实验结果和误差分析组织为一套可复用的课程论文与实验平台。",
        "全文共分六个部分。第一部分为绪论，说明研究背景和问题定义；第二部分介绍 BM25、文本嵌入和向量索引等理论基础；第三部分说明系统架构与工程实现；第四部分给出真实数据上的实验设计与结果分析；第五部分总结误差来源、不足与改进方向；第六部分对全文进行总结和展望。",
    ]:
        add_body_paragraph(doc, text)


def add_section_theory(doc: Document) -> None:
    doc.add_heading("二 相关理论与技术基础", level=1)
    doc.add_heading("2.1 信息检索基本流程", level=2)
    for text in [
        "一个完整的信息检索系统通常包括数据采集、文本清洗、索引构建、查询处理、相关性排序和效果评测六个环节。本文以商品评论为检索对象，以自然语言需求为查询输入，在评论级排序之后再聚合为商品级结果，本质上属于“评论检索驱动商品检索”的两阶段检索流程。",
        "与传统文档检索相比，商品评论具有短文本、噪声高、口语化和强主观性的特点；与普通商品搜索相比，用户目标又往往不是查找名称，而是寻找“某种使用体验是否真实存在”。因此，本文既需要保留传统检索中的可解释关键词能力，也需要引入语义表示来处理需求表达差异。",
    ]:
        add_body_paragraph(doc, text)

    doc.add_heading("2.2 BM25 关键词检索", level=2)
    for text in [
        "BM25 是信息检索领域的经典概率排序模型，综合考虑词项频率、逆文档频率与文档长度归一化，被广泛作为全文检索基线。在商品评论场景中，BM25 的优势在于实现稳定、排序逻辑清晰，并且当查询词与评论表述存在显著重合时，可以快速给出高质量结果。",
        "但 BM25 的局限同样明显。当用户使用“适合学生党”“不打扰室友”“带着不累”这类隐含需求表达时，评论证据往往采用完全不同的措辞，导致 BM25 只能依赖偶然的词项共现，难以稳定覆盖真实需求。",
    ]:
        add_body_paragraph(doc, text)

    doc.add_heading("2.3 文本嵌入与语义表示", level=2)
    for text in [
        "文本嵌入模型通过将自然语言映射为连续向量，使语义相近的文本在向量空间中距离更近。例如“声音小”“安静”“不会吵到室友”在字面上不同，但可以通过向量表示获得较高的语义相似度。这为商品评论中复杂、多样的经验反馈匹配提供了技术基础。",
        "然而，纯向量语义检索并不天然等于更好的商品检索。评论集合中经常存在跨类别的相似场景描述，例如“声音小”既可能指键盘，也可能指耳机、台灯甚至风扇。因此，在垂直商品场景中，语义检索需要与类别约束协同工作，而不是完全脱离结构化信息。",
    ]:
        add_body_paragraph(doc, text)

    doc.add_heading("2.4 FAISS 与向量索引", level=2)
    for text in [
        "FAISS 是面向高维向量相似检索的经典工程工具，支持精确检索与多种近似索引结构。本文当前数据规模约为一万条评论，能够使用精确索引完成语义检索，同时保留后续切换至 IVF 或 HNSW 的工程扩展空间。",
        "在本文系统中，FAISS 主要负责评论向量的存储与 Top-K 相似检索；查询经类别约束后，只在对应品类的评论候选集中进行相似度匹配。这样既能够控制跨类别噪声，也有利于提升后续商品聚合的稳定性。",
    ]:
        add_body_paragraph(doc, text)

    doc.add_heading("2.5 检索评价指标", level=2)
    for text in [
        "本文采用评论级与商品级双层评测机制。评论级指标使用 Precision@K、Recall@K 和 MRR，用于衡量系统是否能找回真正有证据价值的评论；商品级指标同样使用 Precision@K、Recall@K 和 MRR，用于衡量系统是否能将相关商品尽早排到前列。",
        "之所以采用双层评测，是因为商品评论检索的最终目标并不是找到“最像查询的一条评论”，而是帮助用户完成商品决策。一个系统即使评论级相似度不够强，只要能将商品及其证据稳定聚合出来，仍然可能在商品级效果上更符合实际应用目标。",
    ]:
        add_body_paragraph(doc, text)


def add_section_system(doc: Document, inputs: dict[str, object]) -> None:
    reviews = inputs["reviews"]
    raw_summary = inputs["raw_summary"]
    benchmark_dir = inputs["benchmark_dir"]

    doc.add_heading("三 系统设计与实现", level=1)
    doc.add_heading("3.1 总体架构", level=2)
    for text in [
        "本文实现的平台采用统一数据层和三类检索子系统的组织方式。数据层负责评论采集、清洗和标注；传统检索子系统负责 BM25 评论排序与商品聚合；语义检索子系统负责类别识别、需求抽取、向量索引和语义匹配；融合检索子系统对 BM25 与语义得分进行归一化和线性融合，最后输出统一的商品结果与评论证据。",
        "整个工程以命令行脚本为入口，支持从原始数据抓取、清洗、建库、单次查询到批量 benchmark 的完整复现。这种设计满足课程实验对可复现实验流程的要求，也方便后续替换语义模型或扩充新数据。",
    ]:
        add_body_paragraph(doc, text)

    doc.add_heading("3.2 数据采集与清洗", level=2)
    for text in [
        f"实验数据来自小米商城公开评价页，围绕键盘、台灯、耳机、笔记本和充电宝 5 个固定品类进行采集。原始抓取共获得 {raw_summary['collected_total']} 条评论记录，评论包含商品标识、商品名称、品类、评论正文以及来源链接等信息。",
        f"清洗阶段主要完成去空、去重、噪声字符裁剪、字段标准化和检索属性收缩。最终用于检索实验的数据仅保留 review_id、product_id、product_name、category 和 clean_text 五个核心字段，共 {len(reviews)} 条记录，涉及 {reviews['product_id'].nunique()} 个商品。",
        "真实评论数据与模拟数据相比更接近实际场景，但也带来了模板化好评、无意义短句、配件类商品与主商品混杂等问题。这些噪声恰好成为检索系统误差分析的重要来源。",
    ]:
        add_body_paragraph(doc, text)

    doc.add_paragraph("表 3-1 清洗后实验数据规模", style="Heading 3")
    add_table(
        doc,
        [
            ["统计项", "数值"],
            ["评论总数", f"{len(reviews)}"],
            ["商品总数", f"{reviews['product_id'].nunique()}"],
            ["键盘评论数", f"{int((reviews['category'] == '键盘').sum())}"],
            ["台灯评论数", f"{int((reviews['category'] == '台灯').sum())}"],
            ["耳机评论数", f"{int((reviews['category'] == '耳机').sum())}"],
            ["笔记本评论数", f"{int((reviews['category'] == '笔记本').sum())}"],
            ["充电宝评论数", f"{int((reviews['category'] == '充电宝').sum())}"],
        ],
    )

    doc.add_heading("3.3 三类检索系统实现", level=2)
    for text in [
        "传统检索系统基于 rank-bm25 构建评论级 BM25 索引。查询进入系统后先进行基本清洗和分词，再在评论集合中检索 Top-K 评论，并按商品聚合同一商品下的高分评论，作为传统基线系统。",
        "语义检索系统先通过规则词典识别查询中的商品类别，再将需求表达送入文本编码器。评论侧使用向量索引进行 Top-K 相似检索，类别过滤在检索前就限制了候选集合，从而降低跨品类误检。当前平台支持两种语义编码路径：一是本地可离线运行的语义编码器，确保课程项目可直接复现；二是 OpenAI-compatible embedding client，与持久化 embedding cache 联动。本次正式实验使用百炼兼容接口 text-embedding-v4 完成评论向量化和查询编码。",
        "融合检索系统使用评论级结果融合策略。具体做法是先分别获取 BM25 与语义检索的评论排序结果，再对两路分数分别做 Min-Max 归一化，并通过线性融合公式进行加权：fused_score = α × semantic_score + (1 - α) × bm25_score。这样可以在关键词明确时利用 BM25 的稀疏匹配优势，在隐含需求场景下保留语义检索的补充能力。",
    ]:
        add_body_paragraph(doc, text)

    doc.add_heading("3.4 商品聚合与结果展示", level=2)
    for text in [
        "系统最终服务于商品检索，因此评论级命中只是中间过程。本文按照 product_id 对评论结果进行分组，并为每个商品保留若干条高分评论作为证据；商品得分则由这些评论的综合得分决定。最终结果以“商品名称 + 品类 + 综合得分 + 评论证据”的结构返回。",
        "这种设计的直接好处在于解释性。用户不仅可以看到系统推荐了哪款商品，还可以看到推荐依据来自哪些评论片段，从而更容易判断推荐是否可信。这也是评论反馈检索相对于普通商品搜索的关键价值所在。",
    ]:
        add_body_paragraph(doc, text)


def add_section_experiment(doc: Document, inputs: dict[str, object]) -> None:
    summary = inputs["summary"]
    queries = inputs["queries"]
    review_groups = inputs["review_groups"]
    product_groups = inputs["product_groups"]
    rows = benchmark_rows(summary)
    benchmark_dir = inputs["benchmark_dir"]

    doc.add_heading("四 实验设计与结果分析", level=1)
    doc.add_heading("4.1 实验环境与评测设置", level=2)
    for text in [
        "所有实验均在同一台 Windows 环境机器上完成，核心流程包括数据采集、评论清洗、索引构建、单次查询与批量 benchmark。实验采用 Top-10 评测设置，统一比较 BM25、语义检索和融合检索三类系统在评论级与商品级的效果差异。",
        f"查询集共 {len(queries)} 条，其中关键词明确型、体验反馈型和场景化隐含需求型各 5 条；评论级 qrels 共 {len(inputs['review_qrels'])} 条，商品级 qrels 共 {len(inputs['product_qrels'])} 条。这样的设计既能覆盖基础关键词查询，也能覆盖更能体现用户需求理解难点的场景化表达。",
    ]:
        add_body_paragraph(doc, text)

    doc.add_paragraph("表 4-1 实验环境配置", style="Heading 3")
    add_table(
        doc,
        [
            ["项目", "配置"],
            ["操作系统", "Windows"],
            ["编程语言", "Python 3.14"],
            ["传统检索库", "rank-bm25"],
            ["向量检索库", "faiss-cpu"],
            ["数据处理", "pandas / numpy"],
            ["评测方式", "评论级与商品级 Top-10"],
        ],
    )

    doc.add_paragraph("表 4-2 查询集构成", style="Heading 3")
    add_table(
        doc,
        [
            ["查询类型", "数量", "示例"],
            ["关键词明确型", "5", "静音机械键盘 / 护眼台灯"],
            ["体验反馈型", "5", "耳机久戴耳朵会不会疼"],
            ["场景化隐含需求型", "5", "出差带着跑客户希望电脑轻一点续航久一点"],
        ],
    )

    doc.add_heading("4.2 整体实验结果", level=2)
    add_body_paragraph(
        doc,
        "表 4-3 给出了三套系统在评论级与商品级上的整体指标。可以看到，当前使用百炼兼容 embedding 重跑后的结果中，BM25 在评论级结果中仍然保持很强竞争力，而融合检索在评论级 MRR 和商品级排序上继续表现最好。"
    )

    doc.add_paragraph("表 4-3 三套系统整体 benchmark 结果", style="Heading 3")
    add_table(
        doc,
        [
            ["系统", "评测层级", "P@10", "Recall@10", "MRR"],
            ["BM25", "评论级", fmt(rows["bm25_reviews"]["precision_at_k"]), fmt(rows["bm25_reviews"]["recall_at_k"]), fmt(rows["bm25_reviews"]["mrr"])],
            ["BM25", "商品级", fmt(rows["bm25_products"]["precision_at_k"]), fmt(rows["bm25_products"]["recall_at_k"]), fmt(rows["bm25_products"]["mrr"])],
            ["语义检索", "评论级", fmt(rows["semantic_reviews"]["precision_at_k"]), fmt(rows["semantic_reviews"]["recall_at_k"]), fmt(rows["semantic_reviews"]["mrr"])],
            ["语义检索", "商品级", fmt(rows["semantic_products"]["precision_at_k"]), fmt(rows["semantic_products"]["recall_at_k"]), fmt(rows["semantic_products"]["mrr"])],
            ["融合检索", "评论级", fmt(rows["hybrid_reviews"]["precision_at_k"]), fmt(rows["hybrid_reviews"]["recall_at_k"]), fmt(rows["hybrid_reviews"]["mrr"])],
            ["融合检索", "商品级", fmt(rows["hybrid_products"]["precision_at_k"]), fmt(rows["hybrid_products"]["recall_at_k"]), fmt(rows["hybrid_products"]["mrr"])],
        ],
    )

    add_body_paragraph(
        doc,
        "这一结果说明了一个很重要的现实情况：即便替换为商用兼容 embedding，真实商品评论中的模板化评价、配件类商品混杂和短文本噪声仍会显著影响纯语义评论排序，使得评论级 semantic baseline 依然不天然优于 BM25；但当检索目标从单条评论转向“商品 + 证据”时，类别约束与聚合机制能够帮助语义信息发挥作用，最终使融合检索在商品级排序上明显优于传统 BM25。"
    )

    add_figure(
        doc,
        benchmark_dir / "figures" / "mrr.png",
        "图 4-1 三套系统在评论级与商品级上的 MRR 对比",
    )
    add_figure(
        doc,
        benchmark_dir / "figures" / "precision_at_k.png",
        "图 4-2 三套系统 Precision@10 对比",
    )
    add_figure(
        doc,
        benchmark_dir / "figures" / "recall_at_k.png",
        "图 4-3 三套系统 Recall@10 对比",
    )

    doc.add_heading("4.3 不同查询类型下的结果分析", level=2)
    add_body_paragraph(
        doc,
        "为了观察不同查询形态对系统表现的影响，本文进一步按查询类型统计了三套系统在评论级和商品级的平均指标。表 4-4 聚焦评论级结果，表 4-5 聚焦商品级结果。"
    )

    add_table(
        doc,
        [
            ["表 4-4 评论级分类型结果", "", "", ""],
            ["系统/类型", "P@10", "Recall@10", "MRR"],
            ["BM25-关键词明确型", fmt(review_groups["bm25"]["关键词明确型"].precision), fmt(review_groups["bm25"]["关键词明确型"].recall), fmt(review_groups["bm25"]["关键词明确型"].mrr)],
            ["BM25-体验反馈型", fmt(review_groups["bm25"]["体验反馈型"].precision), fmt(review_groups["bm25"]["体验反馈型"].recall), fmt(review_groups["bm25"]["体验反馈型"].mrr)],
            ["BM25-场景化隐含需求型", fmt(review_groups["bm25"]["场景化隐含需求型"].precision), fmt(review_groups["bm25"]["场景化隐含需求型"].recall), fmt(review_groups["bm25"]["场景化隐含需求型"].mrr)],
            ["语义-关键词明确型", fmt(review_groups["semantic"]["关键词明确型"].precision), fmt(review_groups["semantic"]["关键词明确型"].recall), fmt(review_groups["semantic"]["关键词明确型"].mrr)],
            ["语义-体验反馈型", fmt(review_groups["semantic"]["体验反馈型"].precision), fmt(review_groups["semantic"]["体验反馈型"].recall), fmt(review_groups["semantic"]["体验反馈型"].mrr)],
            ["语义-场景化隐含需求型", fmt(review_groups["semantic"]["场景化隐含需求型"].precision), fmt(review_groups["semantic"]["场景化隐含需求型"].recall), fmt(review_groups["semantic"]["场景化隐含需求型"].mrr)],
            ["融合-关键词明确型", fmt(review_groups["hybrid"]["关键词明确型"].precision), fmt(review_groups["hybrid"]["关键词明确型"].recall), fmt(review_groups["hybrid"]["关键词明确型"].mrr)],
            ["融合-体验反馈型", fmt(review_groups["hybrid"]["体验反馈型"].precision), fmt(review_groups["hybrid"]["体验反馈型"].recall), fmt(review_groups["hybrid"]["体验反馈型"].mrr)],
            ["融合-场景化隐含需求型", fmt(review_groups["hybrid"]["场景化隐含需求型"].precision), fmt(review_groups["hybrid"]["场景化隐含需求型"].recall), fmt(review_groups["hybrid"]["场景化隐含需求型"].mrr)],
        ],
    )

    add_table(
        doc,
        [
            ["表 4-5 商品级分类型结果", "", "", ""],
            ["系统/类型", "P@10", "Recall@10", "MRR"],
            ["BM25-关键词明确型", fmt(product_groups["bm25"]["关键词明确型"].precision), fmt(product_groups["bm25"]["关键词明确型"].recall), fmt(product_groups["bm25"]["关键词明确型"].mrr)],
            ["BM25-体验反馈型", fmt(product_groups["bm25"]["体验反馈型"].precision), fmt(product_groups["bm25"]["体验反馈型"].recall), fmt(product_groups["bm25"]["体验反馈型"].mrr)],
            ["BM25-场景化隐含需求型", fmt(product_groups["bm25"]["场景化隐含需求型"].precision), fmt(product_groups["bm25"]["场景化隐含需求型"].recall), fmt(product_groups["bm25"]["场景化隐含需求型"].mrr)],
            ["语义-关键词明确型", fmt(product_groups["semantic"]["关键词明确型"].precision), fmt(product_groups["semantic"]["关键词明确型"].recall), fmt(product_groups["semantic"]["关键词明确型"].mrr)],
            ["语义-体验反馈型", fmt(product_groups["semantic"]["体验反馈型"].precision), fmt(product_groups["semantic"]["体验反馈型"].recall), fmt(product_groups["semantic"]["体验反馈型"].mrr)],
            ["语义-场景化隐含需求型", fmt(product_groups["semantic"]["场景化隐含需求型"].precision), fmt(product_groups["semantic"]["场景化隐含需求型"].recall), fmt(product_groups["semantic"]["场景化隐含需求型"].mrr)],
            ["融合-关键词明确型", fmt(product_groups["hybrid"]["关键词明确型"].precision), fmt(product_groups["hybrid"]["关键词明确型"].recall), fmt(product_groups["hybrid"]["关键词明确型"].mrr)],
            ["融合-体验反馈型", fmt(product_groups["hybrid"]["体验反馈型"].precision), fmt(product_groups["hybrid"]["体验反馈型"].recall), fmt(product_groups["hybrid"]["体验反馈型"].mrr)],
            ["融合-场景化隐含需求型", fmt(product_groups["hybrid"]["场景化隐含需求型"].precision), fmt(product_groups["hybrid"]["场景化隐含需求型"].recall), fmt(product_groups["hybrid"]["场景化隐含需求型"].mrr)],
        ],
    )

    for text in [
        "从评论级结果看，关键词明确型查询仍然最适合 BM25，这说明真实评论中“静音”“护眼”“大容量”等高辨识词在评论文本中具有较强的稀疏匹配价值；百炼 embedding 提升了融合检索的评论级 MRR，但纯语义评论排序仍受短文本噪声影响。",
        f"从商品级结果看，融合检索在三类查询上都更稳定，尤其在场景化隐含需求型查询中，MRR 从 BM25 的 {product_groups['bm25']['场景化隐含需求型'].mrr:.4f} 提升到 {product_groups['hybrid']['场景化隐含需求型'].mrr:.4f}，说明即使评论级语义排序受噪声影响，类别约束和商品聚合仍然能帮助系统把真正相关的商品更早排到前面。",
        "这也解释了本文为什么要同时做评论级和商品级评测：如果只看评论级结果，容易得出“语义检索不如 BM25”的简单结论；但从最终商品决策角度看，融合检索已经明显优于纯 BM25，更符合电商场景的真实目标。",
    ]:
        add_body_paragraph(doc, text)

    doc.add_heading("4.4 典型案例分析", level=2)
    for text in [
        "以“宿舍夜里写论文想要不吵室友的键盘”为例，BM25 在评论级检索中出现了跨品类现象，部分高分结果实际上来自笔记本相关评论；而语义检索在类别过滤之后能够限制在键盘候选集合内，但由于评论噪声较高，仍会受到模板化评价干扰。融合检索综合了类别约束后的语义得分和关键词得分，在商品层面能够更稳定地把键盘相关商品聚合到前列。",
        "以“出差带着跑客户希望电脑轻一点续航久一点”为例，三套系统都能在商品级将 RedmiBook 14 或同系列产品排到前列，但融合检索更容易把包含“轻薄”“续航够用”“日常办公”这类评论证据的商品稳定提前。这说明在笔记本这类评论信息更充分的类别中，商品聚合可以明显增强最终结果的可解释性。",
    ]:
        add_body_paragraph(doc, text)


def add_section_error_analysis(doc: Document, inputs: dict[str, object]) -> None:
    doc.add_heading("五 误差分析、不足与改进方向", level=1)
    doc.add_heading("5.1 误差来源分析", level=2)
    for text in [
        "第一，真实评论噪声较高。数据中存在模板化好评、无意义短句、与使用体验弱相关的物流评价，以及明显无法支撑商品推荐的灌水文本，这会直接削弱评论级语义检索的稳定性。",
        "第二，品类内部商品存在语义混杂。以“键盘”类为例，当前数据里同时包含平板键盘、键盘保护壳和游戏键盘，虽然都属于“键盘”大类，但与用户真实需求的细粒度匹配并不完全一致，容易让语义模型把“平板配件评论”误认为“键盘使用体验评论”。",
        "第三，当前类别识别仍以规则词典为主，对显式类别词识别效果较好，但对类别省略、同义称呼和更复杂的口语表达还不够鲁棒。第四，尽管本文已经使用百炼兼容 embedding 完成正式重跑，但商品场景中的评论噪声并不会仅靠更强语义模型自动消失，说明后续还需要配合更细粒度数据清洗和商品属性约束。",
    ]:
        add_body_paragraph(doc, text)

    doc.add_heading("5.2 本文工作的局限", level=2)
    for text in [
        "本文已经完成从真实评论采集到三套系统评测的完整闭环，但仍存在一些明显边界。首先，数据虽然超过一万条，却仍集中于一个平台和单一品牌生态，品类覆盖虽然符合课程题目要求，但商品品牌的多样性还不够高。",
        "其次，本文的商品聚合策略仍以评论得分平均与证据截取为主，尚未引入销量、价格、评论时间、评分等结构化特征。再次，本文主要围绕文本评论展开，没有把图像、视频和问答内容纳入统一检索范围。",
    ]:
        add_body_paragraph(doc, text)

    doc.add_heading("5.3 改进方向", level=2)
    for text in [
        "后续可以从四个方向继续改进。第一，在当前百炼兼容接口已经跑通的基础上，继续尝试更适合评论检索的商用中文 embedding 或领域模型，并利用当前已完成的 OpenAI-compatible query/runtime 流程直接重跑 benchmark，观察评论级与商品级指标的变化。",
        "第二，引入更细粒度的品类层级和商品属性标签，将“机械键盘”“平板键盘”“蓝牙耳机”“有线耳机”等进一步拆分，减少大类内部混杂带来的语义偏移。",
        "第三，在商品排序阶段融合评分、价格、销量和时间衰减等结构化特征，使检索结果更贴近真实消费决策。第四，扩展到多模态评论检索与评论摘要生成，让系统不仅返回证据，还能自动生成面向用户的需求满足说明。",
    ]:
        add_body_paragraph(doc, text)


def add_section_conclusion(doc: Document) -> None:
    doc.add_heading("六 总结与展望", level=1)
    for text in [
        "本文围绕“面向用户需求理解的商品评论语义检索增强研究”完成了一套真实可运行的课程实验平台。与仅停留在方案设计或模拟数据阶段的实现不同，本文以真实电商评论为基础，完成了数据抓取、字段收缩、检索建库、人工标注、三套系统对比和图表输出。",
        "实验结果表明，BM25 在真实评论中的评论级表现依然强势，说明传统关键词检索并没有过时；但从最终商品排序目标来看，类别约束和语义信息仍然是必要的，尤其在场景化隐含需求查询中，融合检索能够比纯 BM25 更稳定地提前命中相关商品。",
        "这意味着本文的价值不只是“证明语义检索一定优于传统检索”，而是在真实噪声数据上更细致地回答：什么时候传统方法依然有效，什么时候语义信息必须被引入，以及为什么商品级目标需要评论级证据和类别约束共同支撑。未来在完成商用 embedding 正式重跑、丰富多品牌数据之后，这个平台仍可以继续扩展为更完整的商品评论智能检索实验框架。",
    ]:
        add_body_paragraph(doc, text)


def add_references(doc: Document) -> None:
    doc.add_heading("参考文献", level=1)
    references = [
        "[1] 曼宁 C D，拉格万 P，舒策 H. 信息检索导论[M]. 王斌，译. 北京：人民邮电出版社，2010.",
        "[2] Robertson S E, Zaragoza H. The Probabilistic Relevance Framework: BM25 and Beyond[J]. Foundations and Trends in Information Retrieval, 2009, 3(4): 333-389.",
        "[3] Reimers N, Gurevych I. Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks[C]//Proceedings of EMNLP-IJCNLP. 2019: 3982-3992.",
        "[4] Karpukhin V, Oguz B, Min S, et al. Dense Passage Retrieval for Open-Domain Question Answering[C]//Proceedings of EMNLP. 2020: 6769-6781.",
        "[5] Johnson J, Douze M, Jégou H. Billion-scale similarity search with GPUs[J]. IEEE Transactions on Big Data, 2021, 7(3): 535-547.",
        "[6] Malkov Y A, Yashunin D A. Efficient and robust approximate nearest neighbor search using Hierarchical Navigable Small World graphs[J]. IEEE Transactions on Pattern Analysis and Machine Intelligence, 2020, 42(4): 824-836.",
        "[7] Thakur N, Reimers N, Rücklé A, et al. BEIR: A Heterogeneous Benchmark for Zero-shot Evaluation of Information Retrieval Models[C]//NeurIPS Datasets and Benchmarks. 2021.",
        "[8] Facebook AI Research. FAISS: A library for efficient similarity search and clustering of dense vectors[EB/OL].",
        "[9] 王鹏，张华. 面向用户评论的商品特征挖掘与情感分析研究[J]. 情报科学，2021, 39(6): 102-108.",
        "[10] 李明，陈晨. 基于深度语义表示的文本检索方法研究[J]. 数据分析与知识发现，2022, 6(8): 45-54.",
        "[11] 张宇，王健. 基于向量检索的语义搜索系统设计与实现[J]. 计算机工程与应用，2023, 59(12): 120-128.",
        "[12] OpenAI. Embeddings API documentation[EB/OL].",
    ]
    for ref in references:
        add_body_paragraph(doc, ref, first_line_indent=0.0)


def add_acknowledgement(doc: Document) -> None:
    doc.add_heading("致谢", level=1)
    for text in [
        "感谢《信息服务与信息检索》课程的教学内容为本文提供了完整的方法论框架，使我能够将信息采集、索引构建、排序评测与系统实现串联起来，真正完成一次从想法到落地的检索实验。",
        "同时也感谢在项目推进过程中提供讨论和反馈的老师与同学。本文能够从论文设想到真实数据、真实系统和真实 benchmark 结果，离不开一次次对需求、实验和代码细节的反复打磨。",
    ]:
        add_body_paragraph(doc, text)


def add_appendix(doc: Document) -> None:
    doc.add_heading("附录：核心复现实验命令", level=1)
    appendix_blocks = [
        "1. 数据采集与清洗",
        "python scripts/collect/scrape_xiaomi_reviews.py --per-category-target 3000 --output data/raw/xiaomi_reviews.csv --summary-output data/raw/xiaomi_reviews_summary.json",
        "python scripts/preprocess/clean_reviews.py data/raw/xiaomi_reviews.csv data/processed/xiaomi_reviews_clean.csv",
        "python scripts/preprocess/slim_reviews_for_retrieval.py data/processed/xiaomi_reviews_clean.csv data/processed/xiaomi_reviews_retrieval.csv",
        "2. 三套检索系统建库",
        "python scripts/build_index/build_bm25_index.py data/processed/xiaomi_reviews_retrieval.csv outputs/indexes/bm25_xiaomi.json",
        "python scripts/build_index/embed_reviews.py data/processed/xiaomi_reviews_retrieval.csv outputs/semantic/xiaomi_bailian_embeddings.jsonl --cache outputs/semantic/bailian_embedding_cache.jsonl",
        "python scripts/build_index/build_faiss_index.py data/processed/xiaomi_reviews_retrieval.csv outputs/semantic/xiaomi_bailian_embeddings.jsonl outputs/indexes/semantic_xiaomi_bailian.json",
        "3. 批量 benchmark",
        "python scripts/evaluate/run_full_benchmark.py --annotations-dir data/annotations --bm25-index outputs/indexes/bm25_xiaomi.json --semantic-index outputs/indexes/semantic_xiaomi_bailian.json --embedding-cache outputs/semantic/bailian_embedding_cache.jsonl --output-dir outputs/runs/xiaomi_bailian_benchmark --top-k 10 --product-top-n 10 --alpha 0.55",
    ]
    for block in appendix_blocks:
        p = doc.add_paragraph()
        if block.startswith("python "):
            p.paragraph_format.first_line_indent = Cm(0)
            run = p.add_run(block)
            set_font(run, "Consolas", 10)
        else:
            run = p.add_run(block)
            set_font(run, "宋体", 12, bold=block[:2].isdigit())


def add_body_paragraph(doc: Document, text: str, *, first_line_indent: float = 0.74) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(first_line_indent)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, "宋体", 12)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(str(value))
            set_font(run, "宋体", 10.5, bold=row_index == 0)
    doc.add_paragraph()


def add_figure(doc: Document, image_path: Path, caption: str) -> None:
    if image_path.exists():
        doc.add_picture(str(image_path), width=Inches(5.8))
        p = doc.paragraphs[-1]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption)
    set_font(run, "宋体", 11)


def add_spacer(doc: Document, count: int) -> None:
    for _ in range(count):
        doc.add_paragraph()


def set_font(run, name: str, size: int | float, bold: bool = False) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold


def fmt(value: float) -> str:
    return f"{value:.4f}"


def benchmark_rows(summary: pd.DataFrame) -> dict[str, dict[str, float]]:
    rows: dict[str, dict[str, float]] = {}
    for row in summary.itertuples(index=False):
        rows[str(row.system)] = {
            "precision_at_k": float(row.precision_at_k),
            "recall_at_k": float(row.recall_at_k),
            "mrr": float(row.mrr),
        }
    return rows


if __name__ == "__main__":
    build_document(TARGET_DOCX)
