from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

from build_thesis_docx import (
    BENCHMARK_DIR,
    ROOT,
    add_body_paragraph,
    add_figure,
    add_spacer,
    add_table,
    benchmark_rows,
    configure_document,
    fmt,
    load_inputs,
    set_font,
)


TARGET_DOCX = ROOT / "面向用户需求理解的商品评论反馈语义检索增强研究-完成版.docx"


def build_document(output_path: Path) -> None:
    inputs = load_inputs(ROOT)
    doc = Document()
    configure_document(doc)

    add_cover(doc)
    add_abstract(doc, inputs)
    add_chapter_one(doc, inputs)
    add_chapter_two(doc)
    add_chapter_three(doc, inputs)
    add_chapter_four(doc, inputs)
    add_chapter_five(doc, inputs)
    add_chapter_six(doc)
    add_references(doc)
    add_acknowledgement(doc)
    add_appendix(doc)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    temp_output = output_path.with_suffix(".tmp.docx")
    doc.save(temp_output)
    temp_output.replace(output_path)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("学校代码：10022")
    set_font(run, "宋体", 12)

    add_spacer(doc, 2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("信息服务与信息检索")
    set_font(run, "黑体", 18, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("2026 年 4 月 27 日")
    set_font(run, "宋体", 12)

    add_spacer(doc, 3)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("面向用户需求理解的商品评论反馈语义检索增强研究")
    set_font(run, "黑体", 20, bold=True)

    add_spacer(doc, 3)

    for line in [
        "作者：赖稳",
        "学号：220401501",
        "专业：信息管理与信息系统",
        "课程：信息服务与检索",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        set_font(run, "宋体", 14)

    doc.add_page_break()


def add_abstract(doc: Document, inputs: dict[str, object]) -> None:
    reviews: pd.DataFrame = inputs["reviews"]  # type: ignore[assignment]
    queries: pd.DataFrame = inputs["queries"]  # type: ignore[assignment]
    review_qrels: pd.DataFrame = inputs["review_qrels"]  # type: ignore[assignment]
    product_qrels: pd.DataFrame = inputs["product_qrels"]  # type: ignore[assignment]
    summary: pd.DataFrame = inputs["summary"]  # type: ignore[assignment]
    raw_summary: dict[str, object] = inputs["raw_summary"]  # type: ignore[assignment]
    rows = benchmark_rows(summary)

    doc.add_heading("摘要", level=1)
    paragraphs = [
        (
            "随着电子商务平台商品数量与评论规模持续增长，用户在检索商品时越来越依赖评论反馈来理解商品是否真正满足自身需求。"
            "传统商品检索通常依据标题、类目、品牌和参数进行关键词匹配，在属性明确的查询上仍然有效，但对于"
            "“宿舍夜里写论文想要不吵室友的键盘”“出差带着跑客户希望电脑轻一点续航久一点”等场景化、口语化和隐含化需求表达，"
            "标题与参数信息往往不足以支撑用户决策，真正有价值的信息分散在海量评论反馈之中。"
        ),
        (
            "针对这一问题，本文在《面向用户需求理解的商品评论反馈语义检索增强研究》原有结构与论证框架基础上，"
            "完成了从论文设想到真实系统落地的工程实现。系统围绕键盘、台灯、耳机、笔记本和充电宝 5 个固定品类，"
            f"采集小米商城公开评论原始数据 {raw_summary['collected_total']} 条，经过去重、短文本过滤和字段规整后，"
            f"形成 {len(reviews)} 条检索用评论、{reviews['product_id'].nunique()} 个商品实例的实验数据集；"
            f"同时构建了 {len(queries)} 条人工标注查询、{len(review_qrels)} 条评论级相关标注和 {len(product_qrels)} 条商品级相关标注。"
        ),
        (
            "在系统实现上，本文构建了统一的命令行实验平台，并落地了三套可对比的检索系统："
            "基于 BM25 的传统评论检索系统、基于类别约束与向量索引的评论反馈语义检索系统、"
            "以及融合 BM25 分数与语义相似度分数的混合检索系统。语义检索部分接入 OpenAI 兼容嵌入接口调用链路，"
            "在正式实验中采用百炼兼容接口完成评论向量化与查询编码，结合 FAISS 建立评论向量索引，"
            "最终输出“商品结果 + 评论证据”的检索结果形式。"
        ),
        (
            f"正式实验表明，在评论级评测上，BM25 基线取得 Precision@10={rows['bm25_reviews']['precision_at_k']:.4f}、"
            f"Recall@10={rows['bm25_reviews']['recall_at_k']:.4f}、MRR={rows['bm25_reviews']['mrr']:.4f}，"
            "说明真实评论场景中仍然存在较强的词项匹配优势；但在更贴近最终商品决策的商品级评测上，"
            f"融合检索取得了最佳结果，Precision@10={rows['hybrid_products']['precision_at_k']:.4f}、"
            f"Recall@10={rows['hybrid_products']['recall_at_k']:.4f}、MRR={rows['hybrid_products']['mrr']:.4f}。"
            "尤其在场景化隐含需求查询中，类别约束与语义匹配能够更稳定地提前命中相关商品，并通过评论证据解释推荐原因。"
        ),
        (
            "本文不仅验证了商品评论反馈语义检索的可行性，也证明了在真实电商评论数据上，"
            "传统关键词检索、类别约束语义检索与混合检索之间存在清晰的能力边界。"
            "研究结果可为商品搜索、评论证据检索、消费决策辅助和智能信息服务提供可复现的实验平台与实现参考。"
        ),
    ]
    for text in paragraphs:
        add_body_paragraph(doc, text)

    p = doc.add_paragraph()
    run = p.add_run("关键词：")
    set_font(run, "黑体", 12, bold=True)
    run = p.add_run("用户需求理解；商品评论反馈；BM25；语义检索；FAISS；混合检索；信息检索实验")
    set_font(run, "宋体", 12)


def add_chapter_one(doc: Document, inputs: dict[str, object]) -> None:
    doc.add_heading("一 绪论", level=1)

    doc.add_heading("1.1 研究背景", level=2)
    for text in [
        "商品评论是用户真实使用体验的直接记录，也是现代电商信息服务中最具解释力的内容资源之一。"
        "与标题、价格、参数等结构化信息相比，评论文本更容易反映噪声大小、续航表现、便携程度、护眼体验、"
        "佩戴舒适度和宿舍场景适用性等细粒度体验信息。用户在购物决策时，越来越倾向于先从评论中寻找证据，再决定是否购买。",
        "但是，评论检索不同于传统文档检索。用户的查询往往包含一个明确品类与一个隐含需求表达，例如"
        "“适合上课偷偷用的键盘”“租房书桌小想找不占地方又够亮的台灯”“周末露营两天想带一个能给手机和耳机补电的充电宝”。"
        "这类查询的真正目标不是“找到包含同样词语的商品”，而是“找到在真实场景中确实满足该需求的商品及其评论证据”。",
        "在这样的场景下，传统关键词检索依赖词项重合，很容易因为用户表达与评论表达不一致而漏检；"
        "而纯向量检索虽然能捕捉语义相似性，却又可能把同样描述“声音小”“不打扰室友”的不同品类商品混在一起。"
        "因此，围绕“类别识别是否正确、评论证据是否可信、商品结果是否可解释”构建新的检索增强方案，"
        "是商品评论反馈检索从 demo 走向可用系统的关键。"
    ]:
        add_body_paragraph(doc, text)

    doc.add_heading("1.2 研究目的与意义", level=2)
    for text in [
        "本文的直接目标，是在原始论文构想基础上，完成一个可运行、可复现、可对比的商品评论反馈语义检索平台。"
        "平台不只包含单一语义检索模块，而是统一完成数据采集、评论清洗、基线系统、语义系统、融合系统、人工标注和评测出图，"
        "使论文中的实验设计真正具备程序实现与结果支撑。",
        "从理论意义上看，本文将信息检索课程中的数据采集、索引构建、相关性排序和检索评测，与文本嵌入、向量索引和类别约束机制结合起来，"
        "展示了传统检索与语义检索在真实商品评论场景中的互补关系。这一过程也说明，语义检索并不意味着完全放弃结构化约束，"
        "而是应与具体业务中的类目和商品属性共同作用。",
        "从实践意义上看，本文构建的系统能够以“商品结果 + 评论证据”的形式返回结果，既能提升商品搜索体验，"
        "也能为智能客服、评论摘要、消费决策辅助和评论问答等场景提供证据检索基础。"
    ]:
        add_body_paragraph(doc, text)

    doc.add_heading("1.3 研究内容与论文结构", level=2)
    for text in [
        "本文围绕五个方面展开：一是基于真实公开页面完成 10000+ 商品评论的采集、清洗与规整；"
        "二是实现 BM25 传统检索系统，作为论文对比基线；三是实现基于类别约束和 embedding 的评论反馈语义检索系统；"
        "四是实现融合检索系统，对比传统检索与语义检索在商品级结果上的综合表现；五是构建人工标注查询集和双层评测机制，"
        "完成实验、误差分析与论文写作。",
        "全文共分六章。第一章为绪论，介绍研究背景、目的与意义；第二章阐述相关理论与技术基础；"
        "第三章说明系统架构、数据流程和三套检索系统的实现；第四章给出真实实验设计、结果与案例分析；"
        "第五章讨论误差来源、研究不足与后续改进方向；第六章总结全文并展望后续工作。"
    ]:
        add_body_paragraph(doc, text)


def add_chapter_two(doc: Document) -> None:
    doc.add_heading("二 相关理论与技术基础", level=1)

    doc.add_heading("2.1 信息检索基本流程", level=2)
    for text in [
        "信息检索系统通常包含信息采集、文本加工、索引构建、查询处理、相关性排序和结果评测六个核心环节[1]。"
        "本文将检索对象从网页或论文文档切换为商品评论文本，将最终目标从“找到相关文档”扩展为“通过评论证据支持商品推荐”，"
        "因此在传统流程之上增加了类别识别和商品聚合两个针对业务场景的重要步骤。",
        "在本文平台中，评论是底层检索单元，商品是最终排序单元。系统先在评论层面完成相关性匹配，再把高分评论聚合回商品，"
        "从而兼顾评论证据的细粒度表达与商品结果的最终可用性。"
    ]:
        add_body_paragraph(doc, text)

    doc.add_heading("2.2 关键词检索与 BM25 模型", level=2)
    for text in [
        "BM25 是概率相关性框架中的经典排序模型，综合考虑词频、逆文档频率和文档长度归一化，在通用信息检索任务中长期作为高质量基线[2][3]。"
        "在商品评论检索场景中，当查询与评论存在明显词项重合时，BM25 仍能快速返回高相关结果，因此具有工程稳定、实现简单和易于解释的优势。",
        "但 BM25 的弱点同样明显：当用户的需求表达与评论证据之间缺乏直接词面重合时，系统只能依赖偶然共现，"
        "例如“适合学生党”与“价格不贵”“够用”“性价比高”之间的关系难以通过纯关键词稳定建立。"
    ]:
        add_body_paragraph(doc, text)

    doc.add_heading("2.3 文本嵌入与语义表示", level=2)
    for text in [
        "文本嵌入模型通过将自然语言映射为连续向量，使语义相近的句子在向量空间中距离更近。"
        "Sentence-BERT 等工作证明了这种表示在句子级检索任务中的有效性[4]，后续 Dense Retrieval 工作也进一步推动了向量检索在开放域问答中的应用[7]。",
        "在商品评论场景中，语义表示的价值在于捕捉“声音小”“安静”“不影响室友”“适合图书馆”等词面不同但体验相近的表达。"
        "这使系统能够从用户评论中发现比显式属性更细腻的需求证据。"
    ]:
        add_body_paragraph(doc, text)

    doc.add_heading("2.4 向量数据库与 FAISS", level=2)
    for text in [
        "FAISS 是常用的高维向量相似检索工具库，可支持精确检索与多种近似索引结构[5][12]。"
        "HNSW 等图结构索引则进一步提升了大规模近邻检索的效率[6]。本研究当前数据规模约为一万条评论，"
        "因此使用精确向量索引即可完成正式实验，同时在代码接口中保留后续扩展到更大数据规模的可能。",
        "向量索引的引入，使系统能够在评论层面完成 Top-K 相似反馈检索；而类别约束机制则进一步控制候选集范围，"
        "避免因跨品类语义相似导致的误检。"
    ]:
        add_body_paragraph(doc, text)

    doc.add_heading("2.5 检索评测指标", level=2)
    for text in [
        "本文采用评论级与商品级双层评测机制。评论级 Precision@K、Recall@K 和 MRR 用于判断系统能否直接找回相关评论证据；"
        "商品级指标则更贴近用户最终是否能够找到合适商品。这样的设置与 BEIR 等通用检索评测思路一致，即同时考察召回能力和排序质量[8]。",
        "由于论文目标不是单纯比较谁的评论级得分最高，而是比较哪套系统更适合支撑商品检索与消费决策，因此商品级结果在本文中具有更高解释权重。"
    ]:
        add_body_paragraph(doc, text)


def add_chapter_three(doc: Document, inputs: dict[str, object]) -> None:
    reviews: pd.DataFrame = inputs["reviews"]  # type: ignore[assignment]
    raw_summary: dict[str, object] = inputs["raw_summary"]  # type: ignore[assignment]
    category_counts = (
        reviews["category"].value_counts().rename_axis("category").reset_index(name="count")
    )

    doc.add_heading("三 商品评论反馈语义检索系统设计与实现", level=1)

    doc.add_heading("3.1 系统总体架构", level=2)
    for text in [
        "系统整体由四个层次组成：数据层、索引层、检索层和评测层。数据层负责评论采集、清洗和字段规整；"
        "索引层分别构建 BM25 关键词索引与语义向量索引；检索层实现 BM25、类别约束语义检索和融合检索三套系统；"
        "评测层负责读取人工标注查询集，输出评论级与商品级结果、表格和图表。",
        "在程序实现上，系统以命令行脚本组织全流程。`scripts/collect/scrape_xiaomi_reviews.py` 负责真实评论采集，"
        "`scripts/preprocess` 下的脚本负责清洗与字段压缩，`scripts/build_index` 负责建库，"
        "`scripts/run_retrieval` 负责三套系统的查询运行，`scripts/evaluate/run_full_benchmark.py` 负责统一评测与结果输出。"
    ]:
        add_body_paragraph(doc, text)

    doc.add_heading("3.2 评论数据采集与清洗", level=2)
    for text in [
        "考虑到课程项目需要可复现数据与明确字段，本文没有继续依赖体量不足的公开 demo 数据，而是转向真实公开评论页面进行采集。"
        "最终选择小米商城评论页作为主要来源，围绕键盘、台灯、耳机、笔记本和充电宝五类商品建立固定商品目录，再按品类分页抓取评论。",
        f"采集阶段共获取 {raw_summary['collected_total']} 条原始评论。清洗阶段先移除重复评论、极短文本和无意义模板化内容，"
        "再统一保留 review_id、product_id、product_name、category 和 clean_text 五个核心字段，"
        f"形成 {len(reviews)} 条可直接用于建库与评测的检索数据。该设计既满足论文实验需要，也降低了后续索引构建的噪声。"
    ]:
        add_body_paragraph(doc, text)

    table_rows = [["表 3-1 清洗后实验数据集分布", ""], ["商品类别", "评论数量"]]
    for row in category_counts.itertuples(index=False):
        table_rows.append([str(row.category), str(row.count)])
    add_table(doc, table_rows)

    doc.add_heading("3.3 传统检索系统设计", level=2)
    for text in [
        "传统检索系统采用 BM25 作为评论级排序核心。查询输入后，系统先进行基础清洗，再根据 BM25 索引对全部评论进行打分排序，"
        "随后把高分评论按 product_id 聚合为商品结果。该系统的优点在于实现成熟、检索速度快、结果具有较好可解释性，"
        "也是本文进行对比实验的基础基线。",
        "在真实实验中，BM25 的一个典型问题是容易因共享词项而跨品类命中。"
        "例如“宿舍夜里写论文想要不吵室友的键盘”这类查询，单纯依赖关键词时，评论中含有“写论文”“宿舍”“安静”等词的笔记本评论也可能被提前召回。"
    ]:
        add_body_paragraph(doc, text)

    doc.add_heading("3.4 类别约束语义检索系统设计", level=2)
    for text in [
        "语义检索系统首先根据规则词典解析查询中的商品品类与需求表达。`src/common/query_parser.py` 中的 `parse_query` 函数负责抽取品类别名，"
        "将查询切分为 `category` 与 `need_text` 两部分。之后系统仅在对应品类评论集合中进行向量相似检索，降低跨品类误检风险。",
        "评论向量化通过 OpenAI 兼容调用链路完成。`src/semantic_retrieval/embedding_client.py` 中封装了批量 embedding 请求、"
        "缓存命中、失败重试和 batch size 控制机制；正式实验中，系统接入百炼兼容接口并使用 `text-embedding-v4` 模型完成评论编码。"
        "编码后的评论向量以 FAISS 索引组织，实现类别内 Top-K 评论反馈检索。"
    ]:
        add_body_paragraph(doc, text)

    doc.add_heading("3.5 融合检索系统设计", level=2)
    for text in [
        "融合检索系统的目标不是简单叠加两个结果列表，而是在评论级分数上综合关键词匹配优势与语义相似优势。"
        "`src/hybrid_retrieval/fusion.py` 先分别对 BM25 与语义得分做归一化，再以线性权重进行融合，"
        "最后对融合后的评论结果进行商品聚合排序。",
        "该设计适合本研究场景的原因在于：当查询表达较明确时，BM25 得分能够提供强约束；"
        "当用户使用较口语化或场景化表达时，语义相似度又能补足词面不一致带来的缺口。"
        "因此融合检索天然适合作为三套系统中的综合方案。"
    ]:
        add_body_paragraph(doc, text)

    doc.add_heading("3.6 结果组织与可解释性设计", level=2)
    for text in [
        "为了避免系统只返回一个无法解释的商品列表，本文把评论证据作为最终结果的重要组成部分。"
        "每个商品结果都配套展示若干条高分评论，使用户能够直接看到系统为什么推荐该商品。"
        "这一点与原始论文中“评论反馈语义检索增强”的核心思想保持一致，也更符合商品检索的实际使用方式。"
    ]:
        add_body_paragraph(doc, text)


def add_chapter_four(doc: Document, inputs: dict[str, object]) -> None:
    reviews: pd.DataFrame = inputs["reviews"]  # type: ignore[assignment]
    queries: pd.DataFrame = inputs["queries"]  # type: ignore[assignment]
    summary: pd.DataFrame = inputs["summary"]  # type: ignore[assignment]
    review_groups = inputs["review_groups"]  # type: ignore[assignment]
    product_groups = inputs["product_groups"]  # type: ignore[assignment]
    rows = benchmark_rows(summary)

    doc.add_heading("四 系统实验与结果分析", level=1)

    doc.add_heading("4.1 实验环境与实验设计", level=2)
    for text in [
        "本文实验平台基于 Python 构建，评论语义检索部分通过 OpenAI-compatible 接口接入百炼 embedding 服务，"
        "向量索引采用 FAISS，结果输出采用统一 benchmark 流程。正式实验数据集为清洗后的 10495 条评论，"
        "覆盖 37 个商品实例和 5 个固定品类。",
        f"测试查询集共 {len(queries)} 条，平均分布在关键词明确型、体验反馈型和场景化隐含需求型三类。"
        "每条查询同时标注相关评论和相关商品，以支持评论级与商品级双层评测。"
        "这一定义直接服务于本文的核心问题：系统不仅要召回评论，更要帮助用户更早命中真正合适的商品。"
    ]:
        add_body_paragraph(doc, text)

    add_table(
        doc,
        [
            ["表 4-1 三套系统总体评测结果", "", "", ""],
            ["系统", "Precision@10", "Recall@10", "MRR"],
            ["BM25-评论级", fmt(rows["bm25_reviews"]["precision_at_k"]), fmt(rows["bm25_reviews"]["recall_at_k"]), fmt(rows["bm25_reviews"]["mrr"])],
            ["语义-评论级", fmt(rows["semantic_reviews"]["precision_at_k"]), fmt(rows["semantic_reviews"]["recall_at_k"]), fmt(rows["semantic_reviews"]["mrr"])],
            ["融合-评论级", fmt(rows["hybrid_reviews"]["precision_at_k"]), fmt(rows["hybrid_reviews"]["recall_at_k"]), fmt(rows["hybrid_reviews"]["mrr"])],
            ["BM25-商品级", fmt(rows["bm25_products"]["precision_at_k"]), fmt(rows["bm25_products"]["recall_at_k"]), fmt(rows["bm25_products"]["mrr"])],
            ["语义-商品级", fmt(rows["semantic_products"]["precision_at_k"]), fmt(rows["semantic_products"]["recall_at_k"]), fmt(rows["semantic_products"]["mrr"])],
            ["融合-商品级", fmt(rows["hybrid_products"]["precision_at_k"]), fmt(rows["hybrid_products"]["recall_at_k"]), fmt(rows["hybrid_products"]["mrr"])],
        ],
    )

    for text in [
        "表 4-1 表明，在评论级层面，BM25 依然取得了最强的整体表现；这说明真实评论文本中存在较强的词项重合与模板评价，"
        "传统检索并未失效。另一方面，语义检索在评论级上的分数明显更低，反映出真实评论短文本噪声和品类内部商品混杂会显著干扰语义排序。",
        "但如果把目标切换到商品级，就会看到不同结论。融合检索在商品级 MRR 上达到 0.6984，高于 BM25 的 0.5456 和纯语义的 0.6611，"
        "说明即便评论级排序并不完美，类别约束、语义匹配与商品聚合的组合仍然更适合最终商品决策场景。"
    ]:
        add_body_paragraph(doc, text)

    doc.add_heading("4.2 不同查询类型下的结果分析", level=2)
    add_table(
        doc,
        [
            ["表 4-2 评论级分类型结果", "", "", ""],
            ["系统/类型", "P@10", "Recall@10", "MRR"],
            ["BM25-关键词明确型", fmt(review_groups["bm25"]["关键词明确型"].precision), fmt(review_groups["bm25"]["关键词明确型"].recall), fmt(review_groups["bm25"]["关键词明确型"].mrr)],
            ["BM25-体验反馈型", fmt(review_groups["bm25"]["体验反馈型"].precision), fmt(review_groups["bm25"]["体验反馈型"].recall), fmt(review_groups["bm25"]["体验反馈型"].mrr)],
            ["BM25-场景化隐含需求型", fmt(review_groups["bm25"]["场景化隐含需求型"].precision), fmt(review_groups["bm25"]["场景化隐含需求型"].recall), fmt(review_groups["bm25"]["场景化隐含需求型"].mrr)],
            ["语义-关键词明确型", fmt(review_groups["semantic"]["关键词明确型"].precision), fmt(review_groups["semantic"]["关键词明确型"].recall), fmt(review_groups["semantic"]["关键词明确型"].mrr)],
            ["语义-体验反馈型", fmt(review_groups["semantic"]["体验反馈型"].precision), fmt(review_groups["semantic"]["体验反馈型"].recall), fmt(review_groups["semantic"]["体验反馈型"].mrr)],
            ["语义-场景化隐含需求型", fmt(review_groups["semantic"]["场景化隐含需求型"].precision), fmt(review_groups["semantic"]["场景化隐含需求型"].recall), fmt(review_groups["semantic"]["场景化隐含需求型"].mrr)],
            ["融合-关键词明确型", fmt(review_groups["hybrid"]["关键词明确型"].precision), fmt(review_groups["hybrid"]["关键词明确型"].recall), fmt(review_groups["hybrid"]["关键词明确型"].mrr)],
            ["融合-体验反馈型", fmt(review_groups["hybrid"]["体验反馈型"].precision), fmt(review_groups["hybrid"]["体验反馈型"].recall), fmt(review_groups["hybrid"]["体验反馈型"].mrr)],
            ["融合-场景化隐含需求型", fmt(review_groups["hybrid"]["场景化隐含需求型"].precision), fmt(review_groups["hybrid"]["场景化隐含需求型"].recall), fmt(review_groups["hybrid"]["场景化隐含需求型"].mrr)],
        ],
    )

    add_table(
        doc,
        [
            ["表 4-3 商品级分类型结果", "", "", ""],
            ["系统/类型", "P@10", "Recall@10", "MRR"],
            ["BM25-关键词明确型", fmt(product_groups["bm25"]["关键词明确型"].precision), fmt(product_groups["bm25"]["关键词明确型"].recall), fmt(product_groups["bm25"]["关键词明确型"].mrr)],
            ["BM25-体验反馈型", fmt(product_groups["bm25"]["体验反馈型"].precision), fmt(product_groups["bm25"]["体验反馈型"].recall), fmt(product_groups["bm25"]["体验反馈型"].mrr)],
            ["BM25-场景化隐含需求型", fmt(product_groups["bm25"]["场景化隐含需求型"].precision), fmt(product_groups["bm25"]["场景化隐含需求型"].recall), fmt(product_groups["bm25"]["场景化隐含需求型"].mrr)],
            ["语义-关键词明确型", fmt(product_groups["semantic"]["关键词明确型"].precision), fmt(product_groups["semantic"]["关键词明确型"].recall), fmt(product_groups["semantic"]["关键词明确型"].mrr)],
            ["语义-体验反馈型", fmt(product_groups["semantic"]["体验反馈型"].precision), fmt(product_groups["semantic"]["体验反馈型"].recall), fmt(product_groups["semantic"]["体验反馈型"].mrr)],
            ["语义-场景化隐含需求型", fmt(product_groups["semantic"]["场景化隐含需求型"].precision), fmt(product_groups["semantic"]["场景化隐含需求型"].recall), fmt(product_groups["semantic"]["场景化隐含需求型"].mrr)],
            ["融合-关键词明确型", fmt(product_groups["hybrid"]["关键词明确型"].precision), fmt(product_groups["hybrid"]["关键词明确型"].recall), fmt(product_groups["hybrid"]["关键词明确型"].mrr)],
            ["融合-体验反馈型", fmt(product_groups["hybrid"]["体验反馈型"].precision), fmt(product_groups["hybrid"]["体验反馈型"].recall), fmt(product_groups["hybrid"]["体验反馈型"].mrr)],
            ["融合-场景化隐含需求型", fmt(product_groups["hybrid"]["场景化隐含需求型"].precision), fmt(product_groups["hybrid"]["场景化隐含需求型"].recall), fmt(product_groups["hybrid"]["场景化隐含需求型"].mrr)],
        ],
    )

    for text in [
        "从评论级结果看，BM25 在关键词明确型与体验反馈型查询上保持明显优势，说明真实评论中的高频体验词仍然为传统检索提供了强信号。"
        "而纯语义检索虽然能理解需求表达，但容易受到短评噪声和大类内部商品差异影响，导致评论级 Precision@10 偏低。",
        "从商品级结果看，融合检索在三类查询上表现更稳定，尤其在场景化隐含需求型上，"
        f"MRR 从 BM25 的 {product_groups['bm25']['场景化隐含需求型'].mrr:.4f} 提升到 {product_groups['hybrid']['场景化隐含需求型'].mrr:.4f}。"
        "这说明对于“适合图书馆”“不吵室友”“带着不累”这类需要体验解释的查询，"
        "仅靠关键词往往难以提前命中最佳商品，而类别约束与语义信息能够在商品聚合阶段持续发挥作用。"
    ]:
        add_body_paragraph(doc, text)

    doc.add_heading("4.3 图表与典型案例分析", level=2)
    add_figure(doc, BENCHMARK_DIR / "figures" / "mrr.png", "图 4-1 三套系统 MRR 对比图")
    add_figure(doc, BENCHMARK_DIR / "figures" / "precision_at_k.png", "图 4-2 三套系统 Precision@10 对比图")

    for text in [
        "以“宿舍夜里写论文想要不吵室友的键盘”为例，BM25 的高分评论中曾出现与笔记本相关的结果，"
        "说明它会被“写论文”“宿舍”“安静”等共享词项牵引；语义检索在类别过滤后，能够把搜索范围限制在键盘类评论之内，"
        "从而减少跨品类误检；融合检索则在此基础上进一步利用关键词信号强化排序，使最终商品结果更稳定。",
        "再以“出差带着跑客户希望电脑轻一点续航久一点”为例，三套系统都能把 RedmiBook 系列商品排到较前位置，"
        "但融合检索更容易把包含“轻薄”“续航够用”“办公方便”这类评论证据的商品稳定提前。"
        "这说明，当商品评论本身具备较丰富体验表达时，评论反馈检索不仅能找到商品，还能解释为什么该商品适合用户需求。"
    ]:
        add_body_paragraph(doc, text)


def add_chapter_five(doc: Document, inputs: dict[str, object]) -> None:
    doc.add_heading("五 应用价值、误差分析与改进方向", level=1)

    doc.add_heading("5.1 应用价值", level=2)
    for text in [
        "本文系统最直接的应用价值，在于把商品搜索从“参数匹配”提升为“证据驱动的需求匹配”。"
        "用户不再只看到商品列表，而是能看到支持推荐的评论反馈，从而更快判断商品是否适合自身场景。",
        "这种设计对于智能客服、评论证据检索、商品问答、评论摘要生成和消费决策辅助都具有明显价值。"
        "尤其在用户表达需求但说不清参数时，评论反馈检索能够提供比传统商品搜索更强的解释能力。"
    ]:
        add_body_paragraph(doc, text)

    doc.add_heading("5.2 误差来源分析", level=2)
    for text in [
        "第一，真实评论噪声较大。数据中仍然存在模板化好评、物流评价、极短评价和缺乏体验细节的内容，"
        "这会直接削弱语义检索在评论级上的排序能力，也是纯语义评论级指标偏低的重要原因。",
        "第二，当前五个大类内部仍然存在商品子类型混杂。以键盘类为例，平板触控键盘、保护壳键盘和便携键盘虽然都属于键盘大类，"
        "但其使用场景和用户关注点并不完全一致，这会让语义检索把“平板配件评论”误当成“输入设备体验评论”。",
        "第三，现阶段类别识别主要依赖规则词典。该方法对显式品类词识别较稳，但对品类别名、省略表达和更复杂的口语句式仍不够鲁棒。"
        "第四，商品聚合阶段目前主要依赖评论分数，尚未综合价格、销量、评分和时间等结构化特征。"
    ]:
        add_body_paragraph(doc, text)

    doc.add_heading("5.3 本文工作的不足", level=2)
    for text in [
        "本文已经将原始论文中的实验构想落地为真实系统，但仍存在边界。首先，虽然评论规模超过一万条，"
        "数据来源仍集中于单一平台和有限品牌生态，尚不能代表更广泛的电商评论全貌。",
        "其次，本文正式实验已经接入百炼兼容 embedding 接口并重跑 benchmark，"
        "但不同商用模型之间的检索差异还没有在同一套真实标注上展开更系统的横向比较。"
    ]:
        add_body_paragraph(doc, text)

    doc.add_heading("5.4 改进方向", level=2)
    for text in [
        "后续可以从四个方向继续扩展。第一，扩大多平台、多品牌数据来源，并为五个大类建立更细粒度的子类标签；"
        "第二，继续利用当前已经实现的 OpenAI 兼容运行时接口，对不同商用 embedding 模型做正式对比；"
        "第三，在商品聚合阶段融入评分、价格、销量和评论时间等结构化特征；"
        "第四，把评论反馈检索与大模型摘要或问答结合，进一步生成面向用户的需求满足解释。"
    ]:
        add_body_paragraph(doc, text)


def add_chapter_six(doc: Document) -> None:
    doc.add_heading("六 总结与展望", level=1)
    for text in [
        "本文在《面向用户需求理解的商品评论反馈语义检索增强研究》原有论文结构和参考文献体系基础上，"
        "将原先停留在实验构想阶段的内容推进为真实可运行的课程实验平台。通过真实评论采集、字段规整、"
        "三套检索系统实现、人工标注和统一 benchmark，论文中的系统设计、实验方法和结果分析都获得了明确的程序支撑。",
        "实验结果说明，商品评论反馈语义检索并不是简单用向量检索替换传统检索，而是在合适的类别约束和商品聚合机制下，"
        "让语义信息在最终商品决策中发挥作用。BM25 依然在评论级排序上具有价值，融合检索则在商品级目标上表现最稳。"
        "这一结论比“哪种方法绝对更好”更符合真实信息检索系统的工程规律。",
        "未来随着更多真实评论数据、更细粒度商品标签和更强商用 embedding 的引入，"
        "本研究的实验平台仍可继续扩展为更完整的商品评论证据检索与需求理解系统。"
    ]:
        add_body_paragraph(doc, text)


def add_references(doc: Document) -> None:
    doc.add_heading("参考文献", level=1)
    references = [
        "[1] 曼宁 C D，拉格万 P，舒特策 H. 信息检索导论[M]. 王斌，译. 北京：人民邮电出版社，2010.",
        "[2] Robertson S E, Zaragoza H. The Probabilistic Relevance Framework: BM25 and Beyond[J]. Foundations and Trends in Information Retrieval, 2009, 3(4): 333-389.",
        "[3] Robertson S E, Walker S, Jones S, et al. Okapi at TREC-3[C]//Proceedings of the Third Text REtrieval Conference. 1994: 109-126.",
        "[4] Reimers N, Gurevych I. Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks[C]//Proceedings of EMNLP-IJCNLP. 2019: 3982-3992.",
        "[5] Johnson J, Douze M, Jegou H. Billion-scale similarity search with GPUs[J]. IEEE Transactions on Big Data, 2021, 7(3): 535-547.",
        "[6] Malkov Y A, Yashunin D A. Efficient and robust approximate nearest neighbor search using Hierarchical Navigable Small World graphs[J]. IEEE Transactions on Pattern Analysis and Machine Intelligence, 2020, 42(4): 824-836.",
        "[7] Karpukhin V, Oguz B, Min S, et al. Dense Passage Retrieval for Open-Domain Question Answering[C]//Proceedings of EMNLP. 2020: 6769-6781.",
        "[8] Thakur N, Reimers N, Rucklé A, et al. BEIR: A Heterogeneous Benchmark for Zero-shot Evaluation of Information Retrieval Models[C]//NeurIPS Datasets and Benchmarks. 2021.",
        "[9] 王鹏，张华. 面向用户评论的商品特征挖掘与情感分析研究[J]. 情报科学，2021, 39(6): 102-108.",
        "[10] 李明，陈曦. 基于深度语义表示的文本检索方法研究[J]. 数据分析与知识发现，2022, 6(8): 45-54.",
        "[11] 张宇，王俊. 基于向量检索的语义搜索系统设计与实现[J]. 计算机工程与应用，2023, 59(12): 120-128.",
        "[12] Facebook AI Research. FAISS: A library for efficient similarity search and clustering of dense vectors[EB/OL]. 2024.",
        "[13] Python Software Foundation. Python Documentation[EB/OL]. 2024.",
    ]
    for ref in references:
        add_body_paragraph(doc, ref, first_line_indent=0.0)


def add_acknowledgement(doc: Document) -> None:
    doc.add_heading("致谢", level=1)
    for text in [
        "感谢《信息服务与检索》课程提供的信息检索理论框架，使本文能够将数据采集、索引构建、相关性排序和检索评测串联为完整实验。",
        "同时也感谢在选题讨论、系统实现和实验分析过程中提供思路与反馈的老师和同学。"
        "本论文能够从构想到真实系统落地，离不开多轮对需求、数据和实验细节的持续打磨。"
    ]:
        add_body_paragraph(doc, text)


def add_appendix(doc: Document) -> None:
    doc.add_heading("附录：核心代码片段与复现实验命令", level=1)

    add_body_paragraph(doc, "附录选取系统中的关键实现片段，既保留原论文“核心代码示例”的呈现方式，也直接对应本次真实项目中的实际代码文件。")

    add_code_block(
        doc,
        "A. 查询解析与类别识别（src/common/query_parser.py）",
        [
            "from dataclasses import dataclass",
            "import re",
            "",
            "@dataclass(frozen=True)",
            "class ParsedQuery:",
            "    category: str | None",
            "    need_text: str",
            "",
            "def parse_query(text: str) -> ParsedQuery:",
            "    cleaned_text = _clean_text(text)",
            "    match = _find_best_match(cleaned_text)",
            "    if match is None:",
            "        return ParsedQuery(category=None, need_text=cleaned_text)",
            "    category, alias, start = match",
            "    need_text = _clean_text(cleaned_text[:start] + cleaned_text[start + len(alias):])",
            "    return ParsedQuery(category=category, need_text=need_text)",
        ],
    )

    add_code_block(
        doc,
        "B. OpenAI 兼容嵌入请求与重试（src/semantic_retrieval/embedding_client.py）",
        [
            "def build_openai_compatible_embedder(*, api_key, base_url, model, dimensions=None):",
            "    endpoint = base_url.rstrip('/') + '/embeddings'",
            "    def _embed(texts):",
            "        payload = {'model': model, 'input': list(texts)}",
            "        if dimensions is not None:",
            "            payload['dimensions'] = dimensions",
            "        ...",
            "        return retry(_request, attempts=6, delay_seconds=2.0, exceptions=(URLError, TimeoutError))",
            "    return _embed",
        ],
    )

    add_code_block(
        doc,
        "C. 融合检索分数计算（src/hybrid_retrieval/fusion.py）",
        [
            "def fuse_scores(*, bm25_hits, semantic_hits, alpha=0.5):",
            "    normalized_bm25 = normalize_scores(bm25_hits)",
            "    normalized_semantic = normalize_scores(semantic_hits)",
            "    review_ids = set(normalized_bm25) | set(normalized_semantic)",
            "    return {",
            "        review_id: alpha * normalized_semantic.get(review_id, 0.0)",
            "        + (1.0 - alpha) * normalized_bm25.get(review_id, 0.0)",
            "        for review_id in review_ids",
            "    }",
        ],
    )

    add_code_block(
        doc,
        "D. 关键复现实验命令",
        [
            "python scripts/collect/scrape_xiaomi_reviews.py --per-category-target 3000 --output data/raw/xiaomi_reviews.csv --summary-output data/raw/xiaomi_reviews_summary.json",
            "python scripts/preprocess/clean_reviews.py data/raw/xiaomi_reviews.csv data/processed/xiaomi_reviews_clean.csv",
            "python scripts/preprocess/slim_reviews_for_retrieval.py data/processed/xiaomi_reviews_clean.csv data/processed/xiaomi_reviews_retrieval.csv",
            "python scripts/build_index/build_bm25_index.py data/processed/xiaomi_reviews_retrieval.csv outputs/indexes/bm25_xiaomi.json",
            "python scripts/build_index/embed_reviews.py data/processed/xiaomi_reviews_retrieval.csv outputs/semantic/xiaomi_bailian_embeddings.jsonl --cache outputs/semantic/bailian_embedding_cache.jsonl",
            "python scripts/build_index/build_faiss_index.py data/processed/xiaomi_reviews_retrieval.csv outputs/semantic/xiaomi_bailian_embeddings.jsonl outputs/indexes/semantic_xiaomi_bailian.json",
            "python scripts/evaluate/run_full_benchmark.py --annotations-dir data/annotations --bm25-index outputs/indexes/bm25_xiaomi.json --semantic-index outputs/indexes/semantic_xiaomi_bailian.json --embedding-cache outputs/semantic/bailian_embedding_cache.jsonl --output-dir outputs/runs/xiaomi_bailian_benchmark --top-k 10 --product-top-n 10 --alpha 0.55",
        ],
    )


def add_code_block(doc: Document, title: str, lines: list[str]) -> None:
    p = doc.add_paragraph()
    run = p.add_run(title)
    set_font(run, "黑体", 12, bold=True)
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        run = p.add_run(line)
        set_font(run, "Consolas", 10)


if __name__ == "__main__":
    build_document(TARGET_DOCX)
